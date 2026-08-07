#!/usr/bin/env python3
"""
Script para analizar la estructura del Curs de RAMEN
Muestra el contenido para decidir cómo dividirlo
"""

from docx import Document
import os

def read_ramen_document(filepath):
    """Lee el documento del curso de RAMEN"""
    
    if not os.path.exists(filepath):
        print(f"❌ Error: No encontrado {filepath}\n")
        print(f"📝 Asegúrate de que el archivo esté en: {os.getcwd()}")
        return None
    
    print(f"📖 Leyendo: {filepath}\n")
    print("="*70 + "\n")
    
    doc = Document(filepath)
    
    # Mostrar párrafos
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:  # Solo mostrar líneas no vacías
            # Mostrar nivel de encabezado si lo tiene
            if para.style.name.startswith('Heading'):
                print(f"\n{'='*70}")
                print(f"📌 {text} (Heading)")
                print(f"{'='*70}\n")
            else:
                # Limitar líneas muy largas para legibilidad
                if len(text) > 100:
                    print(f"   {text[:100]}...")
                else:
                    print(f"   {text}")
    
    print("\n" + "="*70)
    print(f"\n✅ Lectura completada")
    
    return doc

def analyze_structure(filepath):
    """Analiza la estructura del documento"""
    
    if not os.path.exists(filepath):
        return
    
    doc = Document(filepath)
    
    print("\n📊 ESTRUCTURA DEL DOCUMENTO:\n")
    
    # Buscar secciones principales
    sections = {}
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if para.style.name.startswith('Heading'):
            current_section = text
            sections[current_section] = 0
            print(f"🔹 {text}")
        elif text and current_section:
            sections[current_section] += 1
    
    print(f"\n📈 Resumen:\n")
    for section, count in sections.items():
        print(f"   {section}: {count} líneas")
    
    return sections

if __name__ == "__main__":
    # Cambiar el nombre si es necesario
    ramen_files = [
        "Curs de RAMEN.docx",
        "Curso de RAMEN.docx",
        "RAMEN.docx"
    ]
    
    filepath = None
    for f in ramen_files:
        if os.path.exists(f):
            filepath = f
            break
    
    if not filepath:
        print("❌ No se encontró el archivo de RAMEN")
        print("\n📝 Archivos buscados:")
        for f in ramen_files:
            print(f"   - {f}")
        print(f"\n📂 Carpeta actual: {os.getcwd()}")
    else:
        # Leer y mostrar
        read_ramen_document(filepath)
        
        # Analizar estructura
        analyze_structure(filepath)