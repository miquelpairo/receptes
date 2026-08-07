#!/usr/bin/env python3
"""
Script para dividir Curs de RAMEN en 2 documentos:
1. Caldos (con guía de complementos)
2. Complementos (todos los toppings)
"""

from docx import Document
from docx.shared import Pt
import os

def read_ramen_document(filepath):
    """Lee el documento del RAMEN"""
    doc = Document(filepath)
    
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    return doc, full_text

def extract_sections(doc):
    """Extrae secciones específicas del documento"""
    
    sections = {
        'intro': [],
        'tipos_brou': [],
        'tonkotsu': [],
        'yasai_miso': [],
        'complementos': []
    }
    
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detectar secciones
        if 'Curs de RAMEN' in text and para.style.name.startswith('Heading'):
            current_section = 'intro'
        elif 'TIPUS DE BROU' in text:
            current_section = 'tipos_brou'
        elif 'TONKOTSU RAMEN' in text and 'YASAI' not in text:
            current_section = 'tonkotsu'
        elif 'YASAI MISO RAMEN' in text:
            current_section = 'yasai_miso'
        elif 'Mandonguilles' in text or 'Naruto' in text or 'Ceba xina' in text:
            current_section = 'complementos'
        
        if current_section and text:
            sections[current_section].append(para)
    
    return sections

def create_caldos_document(doc):
    """Crea documento con los caldos"""
    new_doc = Document()
    
    # Título
    title = new_doc.add_heading('RAMEN - CALDOS', level=1)
    
    # Introducción
    new_doc.add_paragraph(
        'En aquest document trobaràs dues recepta de caldos fonamentals per als raàmens. '
        'Cada caldo té una combinació específica de complementos recomanats per obtenir el millor resultat.'
    )
    
    new_doc.add_paragraph()  # Espacio
    
    # SECCIÓN: TIPUS DE BROU
    new_doc.add_heading('Tipus de Brou', level=2)
    
    shoyu_info = """ASSARI (els més clars i lleugers)

• SHOYU (Salsa de soja): Considerada la més japonesa. Es compon de pollastre, al qual se li afegeix un parell de cullerades d'espessa salsa de soja en ser servida. Tradicionalment se li afegeix després rayu (extracte de pebrot vermell amb oli de sèsam) al gust.

• SHIO (Sal): La més simple de totes les varietats, i en la qual se sent la major influència xinesa. Es considera molt popular a Hokkaido. La sopa és transparent i el seu sabor és més directe.

KOTTERI (els més espessos i saborosos)

• TONKOTSU (Ossos de Porc): La sopa està principalment basada en porc, té força cos i contingut gras. El brou generalment és de color blanc. Originari del nord Kyushu.

• MISO: Creat a Sapporo el 1955. Es prepara generalment a base de pollastre, i es barreja amb alguna varietat de miso en ser servida. A Kanto és comú agregar-li una cullerada de mantega just abans de servir."""
    
    new_doc.add_paragraph(shoyu_info)
    
    new_doc.add_paragraph()  # Espacio
    
    # TONKOTSU RAMEN
    new_doc.add_heading('TONKOTSU RAMEN', level=2)
    new_doc.add_heading('(Caldo de ossos de porc - KOTTERI)', level=3)
    
    new_doc.add_paragraph(
        'Aquest és el caldo més ric i aromàtic. Es combina perfectament amb els toppings carnuts com chasu, mandonguilles i ous marinats.'
    )
    
    new_doc.add_heading('Ingredients per 6-8 pax - Brou:', level=3)
    
    ingredients_tonkotsu = [
        '3 litres d\'aigua',
        '2 peus de porc (tallats a la meitat)',
        '1 kg carcasses de pollastre',
        '1 kg ossos de porc (os salat i espinós)',
        '3 Shiitake secs o bolet de temporada',
        '1 tros de gingebre gran (7-8 cm aprox.)',
        '1 manat de ceba xinesa',
        '1 ceba gran o 2 petites',
        '1 porro',
        '1 poma Granny Smith',
        '2 pastanagues',
        '1 cap d\'alls'
    ]
    
    for ing in ingredients_tonkotsu:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    new_doc.add_heading('Elaboració:', level=3)
    
    elaboracio_tonkotsu = """Blanquejar els ossos:
Dins una olla farem bullir els ossos de porc, els peus i les carcasses de pollastre per netejar de brutícia. Quan comenci a bullir, deixem uns 3-4 minuts i seguidament els escorrerem de l'aigua i els rentarem amb aigua freda.

Preparar el brou:
Tornem a ficar els ossos i les carcasses a l'olla juntament amb la ceba xina, ceba pelada, les pastanagues pelades i tallades en 3 o 4 trossos, i la cabeça d'alls tallada per la meitat i sense pelar, el gingebre que no cal pelar-lo, els bolets secs tal qual i la poma tallada a quarts amb pell, cobrim amb aigua.
Arrencar ebullició i deixar bullir a foc lent 4 hores a olla normal (si evapora molt, es pot anar afegint més aigua) o 1 hora 30 minuts amb olla exprés."""
    
    new_doc.add_paragraph(elaboracio_tonkotsu)
    
    new_doc.add_heading('SHOYU TARE (per a Tonkotsu):', level=3)
    
    tare_ingredients = [
        '300 ml de salsa de soja japonesa',
        '150 ml de Sake',
        '150 ml de Mirin',
        '1 tros d\'alga Kombu (uns 5 g)',
        '4 dents d\'all',
        '3 trossets de gingebre de 1 cm',
        '1 tall de ceba xinesa'
    ]
    
    for ing in tare_ingredients:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    elaboracio_tare = """A un bol posem a remull un tros d'alga kombu amb la salsa de soja i deixem reposar almenys 30 min. Si teniu temps fins hi tot podem deixar-ho a la nevera.

Posarem dins una olleta la soja amb l'alga kombu que teníem reservada juntament amb el sake, el mirin i la ceba xinesa tallada a trossos d'un cm, els alls pelats i el gingebre a rodanxes d'un cm sense pelar.

Ho portarem al foc i quan arranquem ebullició apaguem ràpidament, tapem i deixem infusionar uns 15-20 min."""
    
    new_doc.add_paragraph(elaboracio_tare)
    
    new_doc.add_heading('Complementos recomanats per a Tonkotsu:', level=3)
    
    complementos_tonkotsu = """✓ CHASU (Panxeta de porc marinada) - ESSENCIAL
✓ OUS MARINATS - ESSENCIAL
✓ Mandonguilles - Molt recomanat
✓ Ceba xinesa - Recomanat
✓ Alga Nori - Per croquant
✓ Naruto (gambeta en espiral) - Opcional
✓ Ravenets encurtits - Opcional"""
    
    new_doc.add_paragraph(complementos_tonkotsu)
    
    new_doc.add_paragraph()  # Espacio
    
    # YASAI MISO RAMEN
    new_doc.add_heading('YASAI MISO RAMEN', level=2)
    new_doc.add_heading('(Caldo vegetal amb Miso - ASSARI/KOTTERI mixt)', level=3)
    
    new_doc.add_paragraph(
        'Aquesta és la versió vegetal, més lleugera però igual de aromàtica. '
        'Es combina perfectament amb toppings vegetals i la mantega li dona una textura cremosa.'
    )
    
    new_doc.add_heading('Ingredients per 6-8 pax - Brou:', level=3)
    
    ingredients_yasai = [
        '3 litres d\'aigua (pot ser 2 litres d\'aigua i 1 litre dashi vegetal)',
        '1 ceba',
        '2 pastanagues',
        '½ nap daikon',
        '10 unitats shitakes deshidratats',
        '1 manat ceba xinesa',
        '1 porro',
        '1 branca d\'api',
        '1 cap d\'alls',
        '50 g gingebre fresc',
        '6 fulles de col xinesa',
        '20-30 g d\'algues wakame'
    ]
    
    for ing in ingredients_yasai:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    new_doc.add_heading('Elaboració:', level=3)
    
    elaboracio_yasai = """Rentar i pelar totes les verdures i tallar-les a trossos grossos. La cabeça d'alls la deixarem sense pelar i la tallem per la meitat. El gingebre no cal pelar-lo, el posarem tal qual.

Posar totes les verdures dins una olla amb aigua freda abundant juntament amb les algues, el shitake, el cap d'alls i el gingebre. Portar a ebullició (primer a foc fort i tot seguit a foc lent) el brou durant 1 hora i mitja. A l'olla a pressió 45 minuts. Colem i reservem."""
    
    new_doc.add_paragraph(elaboracio_yasai)
    
    new_doc.add_heading('Miso Tare:', level=3)
    
    miso_ingredients = [
        '50 ml de tamari o salsa de soja sense gluten',
        '50 g de miso blanc o vermell'
    ]
    
    for ing in miso_ingredients:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    new_doc.add_paragraph(
        'En un bol de vidre, afegim a parts iguals el miso que vulguem emprar i el tamari o salsa de soja sense gluten. Barregem per tal de fer una pasta homogènia. Reservem.'
    )
    
    new_doc.add_heading('Complementos recomanats per a Yasai Miso:', level=3)
    
    complementos_yasai = """✓ OUS MARINATS - ESSENCIAL
✓ Pak Choi - Molt recomanat (verd, fresc)
✓ Tofu dur - Molt recomanat (proteïna vegetal)
✓ Mantega freda - Essencial (textura cremosa)
✓ Enoki - Recomanat (textura cruixent)
✓ Blat de moro - Recomanat (dulçor)
✓ Llet de soia - Opcional (més cremós)
✓ Ceba xinesa - Recomanat
✓ Alga Nori - Per croquant"""
    
    new_doc.add_paragraph(complementos_yasai)
    
    return new_doc

def create_complementos_document():
    """Crea documento con todos los complementos"""
    new_doc = Document()
    
    # Título
    title = new_doc.add_heading('RAMEN - COMPLEMENTOS I TOPPINGS', level=1)
    
    new_doc.add_paragraph(
        'Aquí trobaràs totes les recepta dels complementos que es combinen amb els caldos. '
        'Cada complement es pot preparar amb anticipació.'
    )
    
    new_doc.add_paragraph()  # Espacio
    
    # CHASU
    new_doc.add_heading('CHASU (Panxeta de Porc Marinada)', level=2)
    new_doc.add_heading('(Per a Tonkotsu - ESSENCIAL)', level=3)
    
    chasu_ingredients = [
        '500 g de panxeta de porc crua (sense pell)',
        '1 rajolí d\'oli d\'oliva o oli neutre',
        '1 litre aigua',
        '300 ml Tare (recepta a la secció de caldos)'
    ]
    
    for ing in chasu_ingredients:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    chasu_elaboracio = """Traurem la crosta a la panxeta i en tallem el tros de panxeta que volem utilitzar.

Dins una paella amb una mica d'oli marquem a foc fort la carn. Cobrim amb l'aigua i deixem bullir amb una tapa una mica oberta uns 45-60 min. Després, afegim el Tare. Deixem coure mínim uns 30 min i la deixem reposar en nevera amb la mateixa salsa durant una nit a ser possible (8-12h). Fins hi tot es pot congelar."""
    
    new_doc.add_paragraph(chasu_elaboracio)
    
    new_doc.add_paragraph()
    
    # OUS MARINATS
    new_doc.add_heading('OUS MARINATS', level=2)
    new_doc.add_heading('(Per a tots els caldos)', level=3)
    
    ous_ingredients = [
        '4 ous M o L (courem algun més per si es trenquen)',
        '300 ml Tare (Shoyu Tare o Miso Tare segons el caldo)'
    ]
    
    for ing in ous_ingredients:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    ous_elaboracio = """Coure l'ou durant 6 min i 15 segons en aigua bullint i refredar amb aigua gel. 

Els pelarem amb molta cura de no trencar-ne cap i els deixarem marinar coberts amb el tare. És més bo d'un dia per un altre (8-12 h a la nevera)."""
    
    new_doc.add_paragraph(ous_elaboracio)
    
    new_doc.add_paragraph()
    
    # MANDONGUILLES
    new_doc.add_heading('Mandonguilles', level=2)
    
    mandong_ingredients = [
        '200 g carn picada de porc',
        '4 dents d\'all',
        '1 trosset de gingebre',
        '3 cebes tendres xineses'
    ]
    
    for ing in mandong_ingredients:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    mandong_elaboracio = """Amb la carn picada fer unes mini mandonguilles barrejant la carn picada amb una mica de gingebre fresc pelat i ratllat, els alls pelats i ratllats i les cebes tendres picades molt menudes. 

Les courem directament en el brou una vegada colat."""
    
    new_doc.add_paragraph(mandong_elaboracio)
    
    new_doc.add_paragraph()
    
    # RAVENETS ENCURTITS
    new_doc.add_heading('Ravenets Encurtits', level=2)
    
    ravenets_ingredients = [
        '4 ravenets rodons',
        '100 ml vinagre d\'arròs',
        '15 g sucre de canya',
        '5 g de sal',
        '5 g d\'alga Kumbu'
    ]
    
    for ing in ravenets_ingredients:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    ravenets_elaboracio = """Rentem amb aigua els ravenets, eliminem la part verda i els tallem el més finets possible. Reservar en un bol. 

Posem una olla petita el vinagre d'arròs, el sucre de canya i la sal. Escalfem remenant fins que es dissolguin els sòlids. Afegim l'alga kombu i tot seguit ho aboquem al bol on tenim els ravenets tallats. 

Deixem que encurteixien almenys 1 hora."""
    
    new_doc.add_paragraph(ravenets_elaboracio)
    
    new_doc.add_paragraph()
    
    # NARUTO
    new_doc.add_heading('Naruto (Gambeta en Espiral)', level=2)
    
    new_doc.add_paragraph(
        'Tallem rodanxes fines (unes 2 per plat) i les reservem fins a l\'hora de muntar el nostre plat de ramen.'
    )
    
    new_doc.add_paragraph()
    
    # CEBA XINESA
    new_doc.add_heading('Ceba Xinesa', level=2)
    
    ceba_ingredients = [
        '4 cebes xineses'
    ]
    
    for ing in ceba_ingredients:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    ceba_elaboracio = """La tallarem el més fina possible amb talls en diagonal i reservarem en aigua freda per suavitzar una mica la fortor i la "baba". 

10 min abans de servir la deixem escorrent a un colador per a que perdi l'excés d'aigua."""
    
    new_doc.add_paragraph(ceba_elaboracio)
    
    new_doc.add_paragraph()
    
    # PAK CHOI
    new_doc.add_heading('Pak Choi (Per a Yasai Miso)', level=2)
    
    pak_ingredients = [
        '1 pak choi'
    ]
    
    for ing in pak_ingredients:
        new_doc.add_paragraph(ing, style='List Bullet')
    
    pak_elaboracio = """Posem dins una olleta aigua a bullir amb una punta de sal. Tallem el pak choi fent una mena de juliana de 2 cm d'amplada. 

Quan l'aigua bulli afegim el pak choi i quan torni a bullir comptarem 30 segons. Escorrem i reservem a un colador."""
    
    new_doc.add_paragraph(pak_elaboracio)
    
    new_doc.add_paragraph()
    
    # TOFU
    new_doc.add_heading('Tofu Dur (Per a Yasai Miso)', level=2)
    
    new_doc.add_paragraph(
        'Tallem uns daus d\'uns 2x2 cm. Reservem amb cura de que no es trenquin. No cal fer-li cap cocció prèvia.'
    )
    
    new_doc.add_paragraph()
    
    # ENOKI
    new_doc.add_heading('Enoki (Per a Yasai Miso)', level=2)
    
    new_doc.add_paragraph(
        'Es un bolet allargat i molt fi. Tallarem la part del peu, rentarem molt ràpidament amb aigua i reservarem fins al moment d\'emplatar.'
    )
    
    new_doc.add_paragraph()
    
    # BLAT DE MORO
    new_doc.add_heading('Blat de Moro (Per a Yasai Miso)', level=2)
    
    new_doc.add_paragraph(
        'Obrim llauna i escorrem molt bé de tot l\'excés de líquid que tinguin. Reservem.'
    )
    
    new_doc.add_paragraph()
    
    # MANTEGA
    new_doc.add_heading('Mantega (Per a Yasai Miso)', level=2)
    
    new_doc.add_paragraph(
        'La mantenim freda a la nevera fins els últims minuts i tallen uns quadrats de 1x1 cm.'
    )
    
    new_doc.add_paragraph()
    
    # LLET DE SOIA
    new_doc.add_heading('Llet de Soia (Per a Yasai Miso - Opcional)', level=2)
    
    new_doc.add_paragraph(
        '500 ml aprox de llet de soia\n\n'
        'Aquest pas es opcional, aconseguirem un ramen més cremós i aromàtic. Una vegada tenim el nostre caldo vegetal colat, afegim la llet de soia al gust.'
    )
    
    new_doc.add_paragraph()
    
    # ALGA NORI
    new_doc.add_heading('Alga Nori', level=2)
    
    new_doc.add_paragraph(
        'A l\'últim moment, tallem una fulla d\'alga nori en 8 trossos i ràpidament servim al bol del nostre ramen.'
    )
    
    new_doc.add_paragraph()
    
    # COCCIÓ FIDEUS
    new_doc.add_heading('Cocció dels Fideus', level=2)
    
    cocció_text = """Sempre ens guiarem per les especificacions del paquet. En el cas de que siguin frescos el temps pot variar entre els 30 segons i els 3 min. Tot depèn del gruix del fideu. 

Una vegada bullits, escorrem bé sense esbandir-los. Aquesta operació s'ha de fer a les acaballes del nostre ramen, sinó, tindrem uns fideus passats de cocció i se'ns poden enganxar.

MUNTATGE FINAL:

1. Al bol on servirem el ramen posarem primer els fideus acabats de coure.
2. Seguidament el nostre brou barrejat amb part del tare.
3. Es l'hora de col·locar els toppings segons el tipus de caldo.
4. Sempre acabant amb la fulla d'alga nori per poder-la menjar el més cruixent possible.

ITADAKIMASU!!"""
    
    new_doc.add_paragraph(cocció_text)
    
    return new_doc

def main():
    input_file = "Curs de RAMEN.docx"
    
    if not os.path.exists(input_file):
        print(f"❌ Error: No encontrado {input_file}")
        return
    
    print("🍜 Dividiendo Curs de RAMEN en 2 documentos...\n")
    
    # Crear documentos
    print("📝 Creando: Recepta Ramen - Caldos.docx")
    caldos_doc = create_caldos_document(Document(input_file))
    caldos_doc.save("Recepta Ramen - Caldos.docx")
    print("   ✓ Guardado")
    
    print("📝 Creando: Recepta Ramen - Complementos.docx")
    complementos_doc = create_complementos_document()
    complementos_doc.save("Recepta Ramen - Complementos.docx")
    print("   ✓ Guardado")
    
    print("\n✅ Completado!")
    print("\nArchivos creados:")
    print("   • Recepta Ramen - Caldos.docx")
    print("   • Recepta Ramen - Complementos.docx")

if __name__ == "__main__":
    main()