#!/usr/bin/env python3
"""
Script para dividir "Les receptes de l'àvia" en recetas individuales
Lee un DOCX con múltiples recetas y crea un DOCX por cada una
"""

from docx import Document
from docx.shared import Pt, RGBColor
import os
from pathlib import Path

def read_avia_document(filepath):
    """Lee el documento de Les receptes de l'àvia"""
    print(f"📖 Leyendo: {filepath}\n")
    doc = Document(filepath)
    
    # Extraer todo el texto
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    return full_text

def find_recipe_titles(text_lines):
    """Encuentra los títulos de las recetas"""
    # Títulos conocidos del índice de l'àvia
    recipe_titles = [
        "Canelons de Nadal (x90)",
        "Peus de porc",
        "Callos",
        "Sípia amb mandonguilles",
        "Mandonguilles de l'Àvia",
        "Estofat",
        "Bunyols de bacallà",
        "Patates americanes",
        "Rostit",
        "Arròs a la cubana",
        'Mongetes amb "almejas"',
        "Macarrons",
        "Bacallà de quaresma",
        "Crema catalana",
        "Pastís de llimona (Magda)",
        "Pasta d'anxoves (Magda)"
    ]
    
    # Encontrar líneas de índice
    recipe_indices = {}
    for i, line in enumerate(text_lines):
        for title in recipe_titles:
            if title.lower() in line.lower():
                recipe_indices[title] = i
                break
    
    return recipe_indices

def divide_recipes(text_lines, recipe_indices):
    """Divide el texto en recetas individuales"""
    recipes = {}
    
    sorted_titles = sorted(recipe_indices.items(), key=lambda x: x[1])
    
    for idx, (title, start_line) in enumerate(sorted_titles):
        # Encontrar el final de la receta
        if idx + 1 < len(sorted_titles):
            end_line = sorted_titles[idx + 1][1]
        else:
            end_line = len(text_lines)
        
        # Extraer contenido de la receta
        recipe_content = text_lines[start_line:end_line]
        
        # Limpiar contenido
        recipe_content = [line for line in recipe_content if line.strip()]
        
        recipes[title] = recipe_content
    
    return recipes

def create_docx_from_content(title, content):
    """Crea un nuevo DOCX con el contenido de una receta"""
    doc = Document()
    
    # Agregar título
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = 1  # Centrado
    
    # Agregar contenido
    for line in content:
        if line.strip():
            doc.add_paragraph(line)
    
    return doc

def save_recipes(recipes, output_folder):
    """Guarda cada receta como un DOCX individual"""
    
    # Crear carpeta de salida si no existe
    Path(output_folder).mkdir(parents=True, exist_ok=True)
    
    print(f"💾 Guardando recetas individuales en: {output_folder}\n")
    
    for title, content in recipes.items():
        # Limpiar nombre para usar como archivo
        filename = f"Recepta {title}.docx"
        filename = filename.replace('(', '').replace(')', '').replace('/', '-')
        filepath = os.path.join(output_folder, filename)
        
        try:
            doc = create_docx_from_content(title, content)
            doc.save(filepath)
            print(f"✓ {filename}")
        except Exception as e:
            print(f"⚠ Error guardando {filename}: {e}")
    
    print(f"\n✅ Recetas guardadas en: {output_folder}")

def main():
    # Configurar rutas
    input_file = "Còpia de Les receptes de l'àvia.docx"  # Cambiar si es necesario
    output_folder = "Receptes_Individuales"
    
    if not os.path.exists(input_file):
        print(f"❌ Error: No encontrado {input_file}")
        print(f"\n📝 Asegúrate de que el archivo esté en la carpeta actual:")
        print(f"   {os.getcwd()}")
        return
    
    # Leer documento
    text_lines = read_avia_document(input_file)
    
    # Encontrar títulos
    print("🔍 Buscando títulos de recetas...\n")
    recipe_indices = find_recipe_titles(text_lines)
    
    if not recipe_indices:
        print("⚠ No se encontraron recetas. Intentando método alternativo...\n")
        recipes = divide_by_manual_list(text_lines)
    else:
        print(f"Encontradas {len(recipe_indices)} recetas:\n")
        for title in recipe_indices.keys():
            print(f"  • {title}")
        
        # Dividir en recetas
        print("\n" + "="*50)
        recipes = divide_recipes(text_lines, recipe_indices)
    
    # Guardar recetas
    print("="*50 + "\n")
    save_recipes(recipes, output_folder)

def divide_by_manual_list(text_lines):
    """Método alternativo: dividir por lista manual de títulos"""
    titles = [
        "Canelons de Nadal",
        "Peus de porc",
        "Callos",
        "Sípia amb mandonguilles",
        "Mandonguilles",
        "Estofat",
        "Bunyols",
        "Patates americanes",
        "Rostit",
        "Arròs",
        "Mongetes",
        "Macarrons",
        "Bacallà",
        "Crema",
        "Pastís",
        "Pasta"
    ]
    
    recipes = {}
    current_recipe = None
    current_content = []
    
    for line in text_lines:
        # Buscar si es un título
        is_title = False
        for title in titles:
            if line.strip().lower().startswith(title.lower()) and len(line.strip()) < 100:
                is_title = True
                # Guardar receta anterior
                if current_recipe:
                    recipes[current_recipe] = current_content
                
                current_recipe = line.strip()
                current_content = [line]
                break
        
        if not is_title and current_recipe:
            current_content.append(line)
    
    # Guardar última receta
    if current_recipe:
        recipes[current_recipe] = current_content
    
    return recipes

if __name__ == "__main__":
    main()