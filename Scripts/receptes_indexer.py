#!/usr/bin/env python3
"""
Script robusto para indexar recetas DOCX + Google Docs
Maneja problemas de importación en Windows
"""

import json
import re
import os
import sys
from pathlib import Path
from docx import Document

print("🍳 Indexador de Recetas\n")

# Intentar importar Google de forma segura
GOOGLE_AVAILABLE = False
try:
    import importlib
    importlib.import_module('google.auth.oauthlib.flow')
    from google.auth.oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    import tempfile
    GOOGLE_AVAILABLE = True
    print("✓ Librerías de Google disponibles\n")
except ImportError as e:
    print(f"⚠️  Librerías de Google no disponibles ({e})")
    print("   Continuando solo con archivos DOCX locales...\n")

RECIPE_FOLDER_ID = "1zVAHBJmyAtamQmFzOAZOqnyfIJqfaPhR"
SINGLE_RECIPE_FILES = ['ramen', 'curs']

def get_drive_service():
    """Obtiene servicio de Google Drive API"""
    if not GOOGLE_AVAILABLE:
        return None
    
    SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
    
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists('credentials.json'):
                print("   ⚠️  No encontrado: credentials.json\n")
                return None
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        
        with open('token.json', 'w') as token:
            token.write(creds.to_json())
    
    return build('drive', 'v3', credentials=creds)

def list_folder_files(service):
    """Lista archivos en Google Drive"""
    try:
        results = service.files().list(
            q=f"'{RECIPE_FOLDER_ID}' in parents and trashed=false",
            spaces='drive',
            fields='files(id, name, mimeType)',
            pageSize=100
        ).execute()
        return results.get('files', [])
    except Exception as e:
        print(f"   Error listando carpeta: {e}")
        return []

def get_google_doc_content(service, doc_id):
    """Lee contenido de Google Doc"""
    try:
        docs_service = build('docs', 'v1', credentials=service._http.request.credentials)
        doc = docs_service.documents().get(documentId=doc_id).execute()
        
        text = ''
        if 'body' in doc and 'content' in doc['body']:
            for elem in doc['body']['content']:
                if 'paragraph' in elem:
                    text += elem['paragraph'].get('rawText', '') + '\n'
                elif 'table' in elem:
                    for row in elem['table']['tableRows']:
                        for cell in row['tableCells']:
                            for content in cell['content']:
                                if 'paragraph' in content:
                                    text += content['paragraph'].get('rawText', '') + ' '
                        text += '\n'
        return text
    except Exception as e:
        print(f"   Error leyendo Google Doc: {e}")
        return ""

def get_docx_files():
    """Lista DOCX locales"""
    return list(Path('.').glob('*.docx'))

def extract_text_from_docx(filepath):
    """Lee DOCX"""
    try:
        doc = Document(filepath)
        text = ''
        for para in doc.paragraphs:
            text += para.text + '\n'
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + ' '
                text += '\n'
        return text
    except Exception as e:
        print(f"   Error: {e}")
        return ""

def is_single_recipe_file(filename):
    """Comprueba si es archivo de receta única"""
    return any(keyword in filename.lower() for keyword in SINGLE_RECIPE_FILES)

def parse_recipes(filename, text, is_single=False):
    """Parsea recetas"""
    recipes = []
    
    if is_single or is_single_recipe_file(filename):
        recipe = parse_single_recipe(text, filename)
        if recipe and recipe.get('name'):
            recipes.append(recipe)
    else:
        recipe_blocks = re.split(
            r'\n\n+(?=[A-Z])|(?=^Ingredients?:)',
            text,
            flags=re.MULTILINE | re.IGNORECASE
        )
        
        for block in recipe_blocks:
            if len(block.strip()) > 20:
                recipe = parse_single_recipe(block, filename)
                if recipe and recipe.get('name'):
                    recipes.append(recipe)
        
        if not recipes:
            recipe = parse_single_recipe(text, filename)
            if recipe and recipe.get('name'):
                recipes.append(recipe)
    
    return recipes

def parse_single_recipe(text, filename):
    """Extrae una receta"""
    if not text or len(text.strip()) < 10:
        return None
    
    recipe = {
        'name': '',
        'source': 'DOCX' if '.docx' in filename.lower() else 'Google Docs',
        'ingredients': [],
        'instructions': '',
        'tags': []
    }
    
    lines = text.split('\n')
    title = None
    for line in lines:
        line = line.strip()
        if line and len(line) > 5 and not line.startswith('-') and 'ingredient' not in line.lower():
            title = line
            break
    
    recipe['name'] = title[:150] if title else filename.replace('.docx', '')
    
    full_text = text.lower()
    
    # Ingredientes
    for pattern in [r'ingredients?:(.+?)(?=preparació|mode|per a|instruccions|$)', r'ingredients?:(.+?)(?=\n\n)']:
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            ingredients_text = match.group(1)
            ingredients = re.findall(
                r'[-•*]\s*(.+?)(?=\n|$)|^\d+\.\s*(.+?)(?=\n|$)',
                ingredients_text,
                re.MULTILINE
            )
            recipe['ingredients'] = [
                (ing[0] if ing[0] else ing[1]).strip()
                for ing in ingredients 
                if (ing[0] or ing[1]).strip() and len((ing[0] or ing[1]).strip()) > 2
            ]
            break
    
    # Instrucciones
    instructions_match = re.search(
        r'(?:preparació|mode|per a|instruccions|pasos|steps)[:\n]*(.+)',
        text,
        re.DOTALL | re.IGNORECASE
    )
    recipe['instructions'] = instructions_match.group(1).strip()[:1500] if instructions_match else text[:500]
    
    # Tags
    tags = set()
    carn = ['porc', 'vedella', 'pollastre', 'xai', 'carn', 'carne']
    peix = ['sípia', 'gambes', 'peix', 'bacallà', 'pescado']
    postres = ['forn', 'dessert', 'pastís', 'crema', 'xocolata', 'pastel']
    verdures = ['patates', 'pastanagues', 'ceba', 'mongetes']
    arros = ['arròs', 'pasta', 'macarrons', 'canelons', 'ramen']
    asian = ['ramen', 'thai', 'tailandès', 'asiàtic', 'hoisin']
    
    if any(w in full_text for w in carn): tags.add('Carn')
    if any(w in full_text for w in peix): tags.add('Peix/Seafood')
    if any(w in full_text for w in postres): tags.add('Postres')
    if any(w in full_text for w in verdures): tags.add('Verdures')
    if any(w in full_text for w in arros): tags.add('Arròs/Pasta')
    if any(w in full_text for w in asian): tags.add('Asiàtic')
    if 'tradicional' in full_text or 'àvia' in full_text: tags.add('Tradicional')
    if 'nadal' in full_text: tags.add('Festiu')
    if any(w in full_text for w in ['ràpid', 'rapido', 'facil']): tags.add('Ràpid')
    if any(w in full_text for w in ['curs', 'tutorial']): tags.add('Tutorial')
    
    recipe['tags'] = list(tags) if tags else ['Altre']
    
    return recipe

# MAIN
all_recipes = []

# DOCX locales
print("📁 Buscando DOCX locales...")
docx_files = get_docx_files()

if docx_files:
    print(f"   {len(docx_files)} archivo(s)\n")
    for docx_file in docx_files:
        print(f"   📄 {docx_file.name}...", end=" ")
        text = extract_text_from_docx(docx_file)
        recipes = parse_recipes(docx_file.name, text, is_single_recipe_file(docx_file.name))
        if recipes:
            all_recipes.extend(recipes)
            print(f"✓ {len(recipes)}")
        else:
            print("⚠")
else:
    print("   (ninguno)\n")

# Google Docs
if GOOGLE_AVAILABLE:
    print("\n☁️  Conectando Google Drive...")
    try:
        service = get_drive_service()
        if service:
            google_files = list_folder_files(service)
            if google_files:
                print(f"   {len(google_files)} archivo(s)\n")
                for gfile in google_files:
                    mime = gfile.get('mimeType', '')
                    if 'document' in mime:
                        print(f"   📄 {gfile['name']}...", end=" ")
                        text = get_google_doc_content(service, gfile['id'])
                        recipes = parse_recipes(gfile['name'], text, is_single_recipe_file(gfile['name']))
                        if recipes:
                            all_recipes.extend(recipes)
                            print(f"✓ {len(recipes)}")
                        else:
                            print("⚠")
            else:
                print("   (ninguno)\n")
    except Exception as e:
        print(f"   Error: {e}\n")

# Guardar
print(f"\n{'='*50}")
print(f"✅ Total: {len(all_recipes)} receptas\n")

with open('receptes_index.json', 'w', encoding='utf-8') as f:
    json.dump({
        'total': len(all_recipes),
        'recipes': all_recipes,
        'generated': True
    }, f, ensure_ascii=False, indent=2)

print("✓ receptes_index.json generado")
print("\n📖 Abre: receptes_buscador.html\n")

if all_recipes:
    print("Receptas indexadas:")
    for recipe in all_recipes[:10]:
        print(f"  • {recipe['name'][:60]}")